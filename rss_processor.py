import os
import re
import xml.etree.ElementTree as ET
import feedparser
from transformers import BartTokenizer, BartForConditionalGeneration
try:
    from transformers import MarianMTModel, MarianTokenizer, pipeline
    HF_AVAILABLE = True
except ImportError:
    HF_AVAILABLE = False
    
try:
    from googletrans import Translator
except ImportError:
    Translator = None

from langdetect import detect, DetectorFactory
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed
import logging
import newspaper
import time
import json
from pathlib import Path
from typing import List, Dict, Optional
import dateutil.parser

from tkinter import Tk, messagebox
from tkinter.filedialog import askopenfilename

# Configuration des logs
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('rss_processor.log'),
        logging.StreamHandler()
    ]
)
logging.getLogger("transformers.modeling_utils").setLevel(logging.ERROR)
logging.getLogger("urllib3.connectionpool").setLevel(logging.WARNING)

DetectorFactory.seed = 0

# --- Classe Config ---
class Config:
    """Configuration centralisée"""
    MAX_ARTICLES_PER_FEED = 3
    MAX_WORKERS = 3  # Réduit pour éviter les rate limits
    REQUEST_TIMEOUT = 10
    RETRY_COUNT = 2
    CACHE_FILE = "translation_cache.json"
    MAX_CONTENT_LENGTH = 1000  # Caractères max pour le résumé
    MIN_SUMMARY_LENGTH = 50
    MAX_SUMMARY_LENGTH = 300

# --- Classe TranslationManager ---
class TranslationManager:
    """Gestionnaire de traduction avec modèles locaux et services externes"""
    
    def __init__(self, cache_file: str = Config.CACHE_FILE):
        self.cache_file = Path(cache_file)
        self.cache = self._load_cache()
        self.request_count = 0
        self.last_request_time = 0
        
        # Modèles de traduction locaux
        self.local_translator = None
        self.marian_models = {}
        
        # Services externes (fallback)
        self.translator = Translator() if Translator else None
        self.session = requests.Session()
        self._setup_session()
        
        # Charger le modèle de traduction local
        self._load_local_translator()

    def _setup_session(self):
        retry_strategy = Retry(
            total=3,
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504],
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        self.session.mount("http://", adapter)
        self.session.mount("https://", adapter)

    def _load_local_translator(self):
        if not HF_AVAILABLE:
            logging.warning("Transformers non disponible, utilisation des services externes")
            return
        try:
            logging.info("Chargement du modèle de traduction local...")
            self.local_translator = pipeline(
                "translation", 
                model="Helsinki-NLP/opus-mt-en-fr",
                device=-1  # CPU
            )
            logging.info("Modèle de traduction local chargé avec succès")
        except Exception as e:
            logging.warning(f"Impossible de charger le modèle local : {e}")
            self.local_translator = None

    def _translate_with_local_model(self, text: str, src_lang: str = "en") -> str:
        if not self.local_translator:
            return None
        try:
            if src_lang not in ['en', 'auto']:
                return None
            max_length = 500
            if len(text) > max_length:
                chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
                translated_chunks = []
                for chunk in chunks:
                    result = self.local_translator(chunk, max_length=512)
                    if result and len(result) > 0:
                        translated_chunks.append(result[0]['translation_text'])
                return ' '.join(translated_chunks)
            else:
                result = self.local_translator(text, max_length=512)
                if result and len(result) > 0:
                    return result[0]['translation_text']
            return None
        except Exception as e:
            logging.warning(f"Erreur modèle local : {e}")
            return None

    def _translate_with_huggingface_api(self, text: str) -> str:
        try:
            url = "https://api-inference.huggingface.co/models/Helsinki-NLP/opus-mt-en-fr"
            headers = {"Content-Type": "application/json"}
            payload = {"inputs": text[:1000]}
            response = self.session.post(url, json=payload, headers=headers, timeout=30)
            if response.status_code == 200:
                result = response.json()
                if isinstance(result, list) and len(result) > 0:
                    translated = result[0].get('translation_text', '')
                    if translated:
                        return translated
            logging.warning(f"Hugging Face API a échoué : {response.status_code}")
            return None
        except Exception as e:
            logging.warning(f"Erreur Hugging Face API : {e}")
            return None

    def _load_cache(self) -> Dict[str, str]:
        if self.cache_file.exists():
            try:
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logging.warning(f"Impossible de charger le cache : {e}")
        return {}

    def _save_cache(self):
        try:
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.cache, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logging.error(f"Impossible de sauvegarder le cache : {e}")

    def _translate_with_mymemory(self, text: str, src_lang: str = "en", dest_lang: str = "fr") -> str:
        try:
            url = "https://api.mymemory.translated.net/get"
            params = {
                "q": text[:500],
                "langpair": f"{src_lang}|{dest_lang}"
            }
            response = self.session.get(url, params=params, timeout=10)
            if response.status_code == 200:
                data = response.json()
                if data.get("responseStatus") == 200:
                    translated_text = data["responseData"]["translatedText"]
                    if translated_text.lower() != text.lower():
                        return translated_text
            logging.warning(f"MyMemory API a échoué pour : {text[:50]}...")
            return None
        except Exception as e:
            logging.warning(f"Erreur MyMemory API : {e}")
            return None

    def _translate_with_libre(self, text: str, src_lang: str = "en", dest_lang: str = "fr") -> str:
        try:
            url = "https://libretranslate.de/translate"
            data = {
                "q": text[:1000],
                "source": src_lang,
                "target": dest_lang,
                "format": "text"
            }
            response = self.session.post(url, data=data, timeout=15)
            if response.status_code == 200:
                result = response.json()
                if "translatedText" in result:
                    return result["translatedText"]
            logging.warning(f"LibreTranslate a échoué pour : {text[:50]}...")
            return None
        except Exception as e:
            logging.warning(f"Erreur LibreTranslate : {e}")
            return None

    def _translate_with_googletrans(self, text: str) -> str:
        if not self.translator:
            return None
        try:
            self._rate_limit()
            translated = self.translator.translate(text, src='auto', dest='fr')
            if translated and hasattr(translated, 'text') and translated.text:
                return translated.text
            return None
        except Exception as e:
            logging.warning(f"Erreur googletrans : {e}")
            return None

    def _rate_limit(self):
        current_time = time.time()
        if current_time - self.last_request_time < 2:
            time.sleep(2)
        self.last_request_time = current_time

    def translate_to_french(self, text: str) -> str:
        if not text or text.strip() == "":
            return text
        text_key = text.strip()[:500]
        if text_key in self.cache:
            return self.cache[text_key]
        src_lang = "en"
        try:
            detected_lang = detect(text)
            if detected_lang == 'fr':
                self.cache[text_key] = text
                return text
            src_lang = detected_lang
        except:
            pass
        translation_methods = [
            lambda: self._translate_with_mymemory(text, src_lang),
            lambda: self._translate_with_libre(text, src_lang),
            lambda: self._translate_with_googletrans(text)
        ]
        for i, method in enumerate(translation_methods):
            try:
                result = method()
                if result and result.strip() and len(result.strip()) > 10:
                    self.cache[text_key] = result
                    self.request_count += 1
                    if self.request_count % 10 == 0:
                        self._save_cache()
                    logging.info(f"Traduction réussie avec méthode {i+1}")
                    return result
            except Exception as e:
                logging.warning(f"Méthode de traduction {i+1} a échoué : {e}")
                continue
        logging.warning(f"Toutes les méthodes de traduction ont échoué pour : {text[:50]}...")
        self.cache[text_key] = text
        return text

    def format_date_french(self, date_str: str) -> str:
        if not date_str:
            return ""
        months_fr = {
            'January': 'janvier', 'February': 'février', 'March': 'mars', 'April': 'avril',
            'May': 'mai', 'June': 'juin', 'July': 'juillet', 'August': 'août',
            'September': 'septembre', 'October': 'octobre', 'November': 'novembre', 'December': 'décembre',
            'Jan': 'janv.', 'Feb': 'févr.', 'Mar': 'mars', 'Apr': 'avr.',
            'May': 'mai', 'Jun': 'juin', 'Jul': 'juill.', 'Aug': 'août',
            'Sep': 'sept.', 'Oct': 'oct.', 'Nov': 'nov.', 'Dec': 'déc.',
            'Monday': 'lundi', 'Tuesday': 'mardi', 'Wednesday': 'mercredi', 'Thursday': 'jeudi',
            'Friday': 'vendredi', 'Saturday': 'samedi', 'Sunday': 'dimanche',
            'Mon': 'lun', 'Tue': 'mar', 'Wed': 'mer', 'Thu': 'jeu',
            'Fri': 'ven', 'Sat': 'sam', 'Sun': 'dim'
        }
        try:
            parsed_date = dateutil.parser.parse(date_str)
            french_date = parsed_date.strftime("%d %B %Y à %H:%M")
            for eng, fr in months_fr.items():
                french_date = french_date.replace(eng, fr)
            return french_date
        except:
            french_date = date_str
            for eng, fr in months_fr.items():
                french_date = french_date.replace(eng, fr)
            return french_date

    def __del__(self):
        self._save_cache()

# --- Classe ArticleSummarizer ---
class ArticleSummarizer:
    def __init__(self):
        self.tokenizer = None
        self.model = None
        self._load_model()

    def _load_model(self):
        try:
            logging.info("Chargement du modèle BART...")
            self.tokenizer = BartTokenizer.from_pretrained("facebook/bart-large-cnn")
            self.model = BartForConditionalGeneration.from_pretrained("facebook/bart-large-cnn")
            logging.info("Modèle BART chargé avec succès")
        except Exception as e:
            logging.error(f"Erreur lors du chargement du modèle : {e}")
            raise

    def summarize_article(self, content: str) -> str:
        if not content or len(content) < 100:
            return content
        try:
            cleaned_content = self._clean_content(content)
            inputs = self.tokenizer(
                [cleaned_content],
                max_length=1024,
                return_tensors="pt",
                truncation=True
            )
            summary_ids = self.model.generate(
                inputs.input_ids,
                num_beams=4,
                min_length=Config.MIN_SUMMARY_LENGTH,
                max_length=Config.MAX_SUMMARY_LENGTH,
                early_stopping=True,
                no_repeat_ngram_size=3
            )
            summary = self.tokenizer.decode(summary_ids[0], skip_special_tokens=True)
            return self._clean_summary(summary)
        except Exception as e:
            logging.error(f"Erreur lors du résumé : {e}")
            return cleaned_content[:Config.MAX_SUMMARY_LENGTH] + "..."

    def _clean_content(self, content: str) -> str:
        content = re.sub(r'<.*?>', '', content)
        content = re.sub(r'\s+', ' ', content)
        return content.strip()[:Config.MAX_CONTENT_LENGTH]

    def _clean_summary(self, summary: str) -> str:
        summary = re.sub(r'<.*?>', '', summary)
        summary = re.sub(r'\s+', ' ', summary)
        return summary.strip()

# --- Classe ArticleExtractor ---
class ArticleExtractor:
    def __init__(self):
        self.session = requests.Session()
        retry_strategy = Retry(
            total=Config.RETRY_COUNT,
            backoff_factor=1,
            status_forcelist=[429, 500, 502, 503, 504],
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        self.session.mount("http://", adapter)
        self.session.mount("https://", adapter)

    def extract_article_content(self, article_url: str, fallback_text: str = "") -> str:
        try:
            article = newspaper.Article(article_url)
            article.download()
            article.parse()
            content = article.text
            if content and len(content) > 100:
                return content
            else:
                logging.info(f"Contenu insuffisant pour {article_url}, utilisation du fallback")
                return fallback_text
        except Exception as e:
            logging.warning(f"Impossible d'extraire l'article {article_url} : {e}")
            return fallback_text

# --- Classe RSSProcessor ---
class RSSProcessor:
    def __init__(self):
        self.translator = TranslationManager()
        self.summarizer = ArticleSummarizer()
        self.extractor = ArticleExtractor()

    def get_articles_from_feed(self, feed_url: str, max_articles: int = Config.MAX_ARTICLES_PER_FEED) -> List[Dict]:
        try:
            feed = feedparser.parse(feed_url)
            if feed.bozo and feed.bozo_exception:
                logging.warning(f"Flux RSS malformé : {feed_url} - {feed.bozo_exception}")
            articles = []
            for entry in feed.entries[:max_articles]:
                date = entry.get('published', entry.get('updated', entry.get('created', '')))
                source = entry.get('source', {}).get('title', feed.feed.get('title', 'Source inconnue'))
                articles.append({
                    'title': entry.title,
                    'link': entry.link,
                    'date': date,
                    'source': source,
                    'summary': entry.get('summary', '')
                })
            logging.info(f"Récupéré {len(articles)} articles de {feed_url}")
            return articles
        except Exception as e:
            logging.error(f"Erreur lors de la récupération du flux {feed_url} : {e}")
            return []

    def process_feed(self, feed: Dict) -> List[Dict]:
        feed_name = feed.get('text', 'Flux sans nom')
        feed_url = feed.get('xmlUrl')
        if not feed_url:
            logging.warning(f"Pas d'URL pour le flux : {feed_name}")
            return []
        logging.info(f"Traitement du flux : {feed_name}")
        articles = self.get_articles_from_feed(feed_url)
        articles_data = []
        for article in articles:
            try:
                content = self.extractor.extract_article_content(article['link'], fallback_text=article.get('summary', ''))
                if not content:
                    logging.info(f"Contenu vide pour l'article : {article['title']}")
                    continue
                title = self.translator.translate_to_french(article['title'])
                if len(content) > 100:
                    content = self.translator.translate_to_french(content)
                summary = self.summarizer.summarize_article(content)
                if summary and len(summary) > 50:
                    try:
                        detected_lang = detect(summary)
                        if detected_lang != 'fr':
                            summary = self.translator.translate_to_french(summary)
                    except:
                        summary = self.translator.translate_to_french(summary)
                french_date = self.translator.format_date_french(article['date'])
                articles_data.append({
                    'title': title,
                    'date': french_date,
                    'source': article['source'],
                    'summary': summary,
                    'link': article['link']
                })
                logging.debug(f"Article traité : {title[:50]}...")
            except Exception as e:
                logging.error(f"Erreur lors du traitement de l'article {article.get('title', 'Sans titre')} : {e}")
                continue
        logging.info(f"Flux {feed_name} : {len(articles_data)} articles traités avec succès")
        return articles_data

    def create_document(self, category_name: str, all_articles: List[Dict]) -> Document:
        doc = Document()
        title = doc.add_heading(f"Revue de presse : {category_name}", level=0)
        self._add_header_footer(doc, category_name)
        doc.add_paragraph("Table des matières")
        doc.add_paragraph("(À générer automatiquement dans Word via Références > Table des matières)")
        doc.add_page_break()
        stats_para = doc.add_paragraph(f"Nombre d'articles : {len(all_articles)}")
        stats_para.add_run(f"\nDate de génération : {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
        doc.add_paragraph("")
        for i, article in enumerate(all_articles, 1):
            heading = doc.add_heading(f"{i}. {article['title']}", level=1)
            meta_para = doc.add_paragraph()
            meta_para.add_run("Date : ").bold = True
            meta_para.add_run(article['date'])
            meta_para.add_run(" | Source : ").bold = True
            meta_para.add_run(article['source'])
            if article.get('link'):
                meta_para.add_run(" | Lien : ").bold = True
                meta_para.add_run(article['link'])
            summary_para = doc.add_paragraph()
            summary_para.add_run("Résumé : ").bold = True
            summary_para.add_run(article['summary'])
            doc.add_paragraph("")
        return doc

    def _add_header_footer(self, doc: Document, category_name: str):
        section = doc.sections[0]
        header = section.header
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_paragraph.text = f"Revue de presse - {category_name}"
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header_paragraph.runs:
            header_paragraph.runs[0].font.size = Pt(12)
            header_paragraph.runs[0].font.bold = True
        footer = section.footer
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_paragraph.text = f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer_paragraph.runs:
            footer_paragraph.runs[0].font.size = Pt(10)
            footer_paragraph.runs[0].font.italic = True

    def process_opml(self, opml_file: str):
        opml_path = Path(opml_file)
        if not opml_path.exists():
            logging.error(f"Fichier OPML non trouvé : {opml_file}")
            return
        try:
            with open(opml_path, 'r', encoding='utf-8') as f:
                opml_content = f.read()
            root = ET.ElementTree(ET.fromstring(opml_content)).getroot()
            for category in root.findall('body/outline'):
                category_name = category.get('text', 'Catégorie')
                category_path = Path(category_name)
                category_path.mkdir(exist_ok=True)
                timestamp = datetime.now().strftime("%Y%m%d_%Hh%M")
                rss_feeds = list(category.findall('outline'))
                logging.info(f"Traitement de {len(rss_feeds)} flux pour '{category_name}'")
                all_articles = []
                with ThreadPoolExecutor(max_workers=Config.MAX_WORKERS) as executor:
                    future_to_feed = {executor.submit(self.process_feed, feed): feed for feed in rss_feeds}
                    for future in tqdm(as_completed(future_to_feed), total=len(rss_feeds), desc=f"{category_name}"):
                        try:
                            articles = future.result(timeout=30)
                            all_articles.extend(articles)
                        except Exception as e:
                            feed = future_to_feed[future]
                            logging.error(f"Erreur pour le flux {feed.get('text', 'inconnu')} : {e}")
                if all_articles:
                    all_articles.sort(key=lambda x: x.get('date', ''), reverse=True)
                    doc = self.create_document(category_name, all_articles)
                    doc_file = category_path / f"{category_name}_{timestamp}.docx"
                    doc.save(str(doc_file))
                    logging.info(f"Document créé : {doc_file} ({len(all_articles)} articles)")
                else:
                    logging.warning(f"Aucun article récupéré pour la catégorie '{category_name}'")
        except Exception as e:
            logging.error(f"Erreur lors du traitement du fichier OPML : {e}")
            raise

# --- Fonction principale avec interface Tkinter ---
def main():
    try:
        Tk().withdraw()
        opml_file = askopenfilename(
            title="Sélectionnez le fichier OPML",
            filetypes=[("OPML files", "*.opml")]
        )
        if not opml_file:
            messagebox.showinfo("Annulé", "Aucun fichier sélectionné. Le programme va se fermer.")
            return
        processor = RSSProcessor()
        processor.process_opml(opml_file)
        messagebox.showinfo("Terminé", "Traitement terminé ! Les documents Word sont créés.")
    except Exception as e:
        messagebox.showerror("Erreur", f"Erreur fatale : {e}")
        logging.error(f"Erreur fatale : {e}")

if __name__ == "__main__":
    main()
